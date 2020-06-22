# -*- coding: utf-8 -*-
"""
Created on Mon Jun 22 13:46:42 2020
@author: Yenny
"""
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

document = Document()
styles = document.styles
 
#生成所有表样式
for s in styles:
    if s.type == WD_STYLE_TYPE.TABLE:
        document.add_paragraph("表格样式 :  "+ s.name)
        table = document.add_table(3,3, style = s)
        heading_cells = table.rows[0].cells
        heading_cells[0].text = '第一列内容'
        heading_cells[1].text = '第二列内容'
        heading_cells[2].text = '第三列内容'
        document.add_paragraph("\n")
 
document.save('Word所有表格样式.docx')
