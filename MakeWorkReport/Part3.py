# -*- coding: utf-8 -*-
"""
Created on Mon Jun 22 14:12:39 2020
@author: Yenny
"""

from faker import Faker
import random
import pandas as pd
import matplotlib.pyplot as plt
from matplotlib.font_manager import FontProperties
from docx import Document
from docx.oxml.ns import qn


# 伪造数据
fake = Faker('zh_CN')  
name = []
score1 = []
score2 = []
score3 = []
number = range(1,21)
for _ in range(20):
    name.append(fake.name())
    score1.append(random.randint(60,100))
    score2.append(random.randint(60,100))
    score3.append(random.randint(60,100))

# 写入Excel
df = pd.DataFrame({
        '学号':number,
        '姓名':name,        
        '语文':score1,
        '数学':score2,
        '英语':score3
        })

df = df.set_index('学号')  
df.to_excel('Part3_学生成绩单.xlsx')

# 读取数据
students = pd.read_excel('Part3_学生成绩单.xlsx')

# 排序名单
students['总分'] = students.语文 + students.数学 + students.英语
students.sort_values(by='总分', inplace=True, ascending=False)
students.reset_index(drop=True, inplace=True)

# 设置字体
font = FontProperties(fname='C:\\Windows\\Fonts\\simfang.ttf', size=16)

# 绘制图表
plt.rcParams['font.sans-serif']=['SimHei'] # 解决图例中文乱码
plt.rcParams['axes.unicode_minus']=False
ax = students.plot.bar(x='姓名', y=['语文','数学','英语'], stacked=True)
plt.title('学生成绩汇总图', fontproperties=font, fontsize=16)
plt.xlabel('姓名', fontproperties=font, fontsize=10)
plt.xticks(fontproperties=font, rotation='45', fontsize=8)
ax.spines['top'].set_visible(False)
ax.spines['right'].set_visible(False)
plt.tight_layout()
plt.savefig('Part3_data.jpg')

# 操作Word
document = Document()
document.add_heading('学生成绩分析报告', level=0)
first_student = students.iloc[0,:]['姓名']
first_score = students.iloc[0,:]['总分']
# 设置格式
document.styles['Normal'].font.name = 'Times New Roman'
document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

p = document.add_paragraph('本次测评，全班共有{}名同学参加考试，其中分数总分排名第一的同学是'.format(len(students.姓名)))
p.add_run(str(first_student)).bold = True
p.add_run('，分数为')
p.add_run(str(first_score)).bold = True
p.add_run('。学生考试总体成绩如下')

table = document.add_table(rows=len(students.姓名)+1, cols=5, style='Medium Shading 1 Accent 5')
table.cell(0,0).text = '姓名'
table.cell(0,1).text = '语文'
table.cell(0,2).text = '数学'
table.cell(0,3).text = '英语'
table.cell(0,4).text = '总分'

for i,(index,row) in enumerate(students.iterrows()):
    table.cell(i+1, 0).text = str(row['姓名'])
    table.cell(i+1, 1).text = str(row['语文'])
    table.cell(i+1, 2).text = str(row['数学'])
    table.cell(i+1, 3).text = str(row['英语'])
    table.cell(i+1, 4).text = str(row['总分'])

document.add_picture('Part3_data.jpg')
document.save('Part3_学生成绩分析报告.docx')
