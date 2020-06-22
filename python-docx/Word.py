# -*- coding: utf-8 -*-
"""
Created on Fri Jun 19 11:18:59 2020
@author: Yenny
"""

from docx import Document
from docx.shared import Inches

# 新建一个文档
document = Document()

# 写入标题
heading0 = document.add_heading('文档标题', 0)
heading1 = document.add_heading('一级标题', level=1)
heading2 = document.add_heading('二级标题', level=2)
heading3 = document.add_heading('三级标题', level=3)

# 写正文
paragraph1 = document.add_paragraph('这是一个段落')
paragraph2 = heading2.insert_paragraph_before('在指定位置之前，插入一个段落')

# 设置字体格式
paragraph0 = heading1.insert_paragraph_before('进行字体设置，部分字体')
paragraph0.add_run('加粗').bold = True
paragraph0.add_run('，部分字体')
paragraph0.add_run('倾斜').italic = True

# 设置段落格式
document.add_paragraph('引用段落', style='Intense Quote')
document.add_paragraph('插入项目符号', style='List Bullet')
document.add_paragraph('插入编号', style='List Number')

# 插入图片
document.add_picture('picture.jpg', width=Inches(5))

# 插入分页符
document.add_page_break() 

# 插入表格
records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3, style="Light Shading Accent 1")
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_paragraph('更换表格样式')

table2 = document.add_table(rows=1, cols=3, style="Light List Accent 5")
hdr_cells = table2.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table2.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc
    
# 保存
document.save('Example.docx')